"""Extract plain text from a resume / job-description file.

Supported file types:
  .pdf            - via pypdf
  .docx           - via python-docx
  .txt / .md      - read as utf-8

PDF/DOCX dependencies are imported lazily so they don't slow down the
app's cold start when nobody actually drops one of those file types.
"""
from __future__ import annotations

from pathlib import Path

SUPPORTED_SUFFIXES: tuple[str, ...] = (".pdf", ".docx", ".txt", ".md")


class UnsupportedFile(ValueError):
    pass


def extract_text(path: str | Path) -> str:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(str(p))
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(p)
    if suffix == ".docx":
        return _from_docx(p)
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="ignore").strip()
    raise UnsupportedFile(
        f"Unsupported file type {suffix!r}. "
        f"Supported: {', '.join(SUPPORTED_SUFFIXES)}"
    )


def _from_pdf(p: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            "PDF import needs `pypdf`. Run:  pip install pypdf"
        ) from e
    reader = PdfReader(str(p))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            # Some malformed pages throw; skip them rather than failing
            # the whole import.
            continue
    return _normalize("\n\n".join(parts))


def _from_docx(p: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "DOCX import needs `python-docx`. Run:  pip install python-docx"
        ) from e
    doc = Document(str(p))
    parts: list[str] = [par.text for par in doc.paragraphs if par.text.strip()]
    # Resumes commonly use tables for skills / experience grids; pull
    # those too.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)
    return _normalize("\n".join(parts))


def _normalize(text: str) -> str:
    """Collapse runs of blank lines and trim."""
    lines = [ln.rstrip() for ln in text.splitlines()]
    out: list[str] = []
    blank = 0
    for ln in lines:
        if ln.strip() == "":
            blank += 1
            if blank <= 1:
                out.append("")
        else:
            blank = 0
            out.append(ln)
    return "\n".join(out).strip()
